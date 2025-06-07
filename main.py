from flask import Flask, request, jsonify
import os
import pandas as pd
import numpy as np
from transformers import pipeline, AutoModelForSequenceClassification, AutoTokenizer
from utils.data_utils import analyze_dataset
from utils.context_utils import detect_context_with_clustering
import requests
from flask_cors import CORS
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import re
import json
import base64
from io import BytesIO
from docx import Document

app = Flask(__name__)
CORS(app, resources={
    r"/extract_insight": {"origins": "http://localhost:3000"},
    r"/process_dataset": {"origins": "http://localhost:3000"},
    r"/analyze_data": {"origins": "http://localhost:3000"},
    r"/clean_data": {"origins": "http://localhost:3000"},
    r"/column_mapping": {"origins": "http://localhost:3000"},
    r"/perform_mapping": {"origins": "http://localhost:3000"},
    r"/report_gen": {"origins": "http://localhost:3000"},
    r"/remove_outlier": {"origins": "http://localhost:3000"},
    r"/chat_bot": {"origins": "http://localhost:3000"}
})

# Increase the file size limit to 1 GB
app.config['MAX_CONTENT_LENGTH'] = 1024 * 1024 * 1024  # 1 GB

# Define directories and paths
data_dir = os.path.join("project_root", "data")
output_dir = os.path.join("project_root", "outputs")
mapped_dir = os.path.join("project_root", "mapped")
os.makedirs(data_dir, exist_ok=True)
os.makedirs(output_dir, exist_ok=True)

# Retry mechanism for loading the model
def load_model_with_retry(model_name, retries=3, timeout=1200):
    for attempt in range(retries):
        try:
            print(f"Attempt {attempt + 1} to load model...")
            model = AutoModelForSequenceClassification.from_pretrained(model_name, use_auth_token=True, timeout=timeout)
            tokenizer = AutoTokenizer.from_pretrained(model_name, use_auth_token=True)
            return model, tokenizer
        except (requests.exceptions.RequestException, Exception) as e:
            print(f"Attempt {attempt + 1} failed: {str(e)}")
            if attempt < retries - 1:
                print("Retrying...")
            else:
                raise Exception(f"Failed to load model after {retries} attempts. Error: {str(e)}")

# Data cleaning and analysis functions
def clean_column_types(df):
    """Clean and convert column types where possible"""
    for col in df.columns:
        if df[col].dtype == 'object':
            df[col] = df[col].astype(str).str.replace(r'[$%,]', '', regex=True)
            df[col] = pd.to_numeric(df[col], errors='ignore')
    return df

def handle_datetime_columns(df):
    """Detect and convert datetime columns"""
    for col in df.columns:
        if df[col].dtype == 'object':
            try:
                converted_col = pd.to_datetime(df[col], errors='coerce')
                if converted_col.notna().sum() > (0.5 * len(df[col])):
                    df[col] = converted_col
            except Exception:
                pass
    return df

def detect_outliers(df):
    """Detect outliers in numeric columns and provide suggestions"""
    outlier_info = {}
    
    for col in df.select_dtypes(include=[np.number]).columns:
        Q1 = df[col].quantile(0.25)
        Q3 = df[col].quantile(0.75)
        IQR = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        outliers = df[(df[col] < lower_bound) | (df[col] > upper_bound)][col]
        
        if not outliers.empty:
            # Create boxplot for visualization
            plt.figure(figsize=(8, 4))
            sns.boxplot(x=df[col])
            plt.title(f"Boxplot for {col}")
            plt.tight_layout()
            
            # Convert plot to base64 for sending to frontend
            buffer = BytesIO()
            plt.savefig(buffer, format='png')
            buffer.seek(0)
            plot_data = base64.b64encode(buffer.getvalue()).decode('utf-8')
            plt.close()
            
            outlier_info[col] = {
                'count': len(outliers),
                'plot': plot_data,
                'min': float(lower_bound),
                'max': float(upper_bound),
                'mean': float(df[col].mean()),
                'median': float(df[col].median())
            }
    
    return outlier_info

def suggest_null_handling(df):
    """Analyze null values and provide handling suggestions"""
    null_info = {}
    
    for col in df.columns:
        null_count = df[col].isnull().sum()
        if null_count > 0:
            percentage = (null_count / len(df)) * 100
            
            # Determine appropriate methods based on data type
            methods = []
            if percentage > 50:
                methods.append("drop_column")
            else:
                methods.append("drop_rows")
            
            if df[col].dtype == 'object':
                methods.append("fill_na")
                methods.append("fill_mode")
            else:
                methods.append("fill_mean")
                methods.append("fill_median")
                methods.append("fill_mode")
            
            null_info[col] = {
                'count': int(null_count),
                'percentage': float(percentage),
                'dtype': str(df[col].dtype),
                'methods': methods,
                'mean': float(df[col].mean()) if df[col].dtype != 'object' else None,
                'median': float(df[col].median()) if df[col].dtype != 'object' else None,
                'mode': str(df[col].mode()[0]) if not df[col].mode().empty else None
            }
    
    return null_info

def clean_data_with_options(df, cleaning_options):
    """Apply cleaning operations based on user-selected options"""
    # Handle outliers
    if 'outliers' in cleaning_options:
        for col, action in cleaning_options['outliers'].items():
            if col in df.columns and df[col].dtype != 'object':
                Q1 = df[col].quantile(0.25)
                Q3 = df[col].quantile(0.75)
                IQR = Q3 - Q1
                lower_bound = Q1 - 1.5 * IQR
                upper_bound = Q3 + 1.5 * IQR
                
                if action == 'remove':
                    df = df[(df[col] >= lower_bound) & (df[col] <= upper_bound)]
                elif action == 'mean':
                    df[col] = df[col].mask((df[col] < lower_bound) | (df[col] > upper_bound), df[col].mean())
                elif action == 'median':
                    df[col] = df[col].mask((df[col] < lower_bound) | (df[col] > upper_bound), df[col].median())
    
    # Handle null values
    if 'nulls' in cleaning_options:
        for col, action in cleaning_options['nulls'].items():
            if col in df.columns:
                if action == 'drop_column':
                    df = df.drop(columns=[col])
                elif action == 'drop_rows':
                    df = df.dropna(subset=[col])
                elif action == 'fill_na' and df[col].dtype == 'object':
                    df[col] = df[col].fillna('N/A')
                elif action == 'fill_mean' and df[col].dtype != 'object':
                    df[col] = df[col].fillna(df[col].mean())
                elif action == 'fill_median' and df[col].dtype != 'object':
                    df[col] = df[col].fillna(df[col].median())
                elif action == 'fill_mode':
                    if not df[col].mode().empty:
                        df[col] = df[col].fillna(df[col].mode()[0])
    
    # Always perform basic cleaning
    df = clean_column_types(df)
    df = handle_datetime_columns(df)
    
    return df

@app.route('/process_dataset', methods=['POST'])
def process_dataset():
    """Process the dataset uploaded by the frontend and analyze its context."""
    # Check if a file is included in the request
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files['file']
    user_id = request.form.get('user_id')
    print(f"user_id: {user_id}")
    if not user_id:
        return jsonify({"error": "User ID is missing"}), 400

    # Check if a file is selected
    if file.filename == '':
        return jsonify({"error": "No file selected for uploading."}), 400

    # Save the uploaded file to the data directory
    try:
        file_name = f"{file.filename}"
        file_path = os.path.join(data_dir, file_name)
        file.save(file_path)


        # Analyze dataset
        df = analyze_dataset(file_path, output_dir)

        # Perform context detection
        report_path = os.path.join(output_dir,"dataset_context_report.txt")

        _, file_extension = os.path.splitext(file_path)
        
        # Load data based on file type
        if file_extension == '.csv':
            df = pd.read_csv(file_path)
        elif file_extension in ['.xls', '.xlsx']:
            df = pd.read_excel(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Please provide CSV or Excel file.")
        detect_context_with_clustering(df, report_path)

        return jsonify({
            "message": "Dataset uploaded and processed successfully!", 
            "context_report": report_path
        }), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/extract_insight', methods=['POST'])
def extract_insight():
    """Extract business insight based on user description."""
    user_description = request.json.get('user_description', None)
    user_id = request.form.get('user_id')
    try:
        # Get the uploaded dataset from the data directory
        dataset_files = [f for f in os.listdir(data_dir) if (f.endswith('.xlsx') or f.endswith('.csv'))]
        if not dataset_files:
            return jsonify({"error": "No dataset found in the /data/ directory."}), 400

        dataset_path = os.path.join(data_dir, dataset_files[0])
        if not os.path.exists(dataset_path):
            raise FileNotFoundError(f"Dataset path {dataset_path} does not exist.")
        
        _, file_extension = os.path.splitext(dataset_path)
        
        # Load data based on file type
        if file_extension == '.csv':
            df = pd.read_csv(dataset_path)
        elif file_extension in ['.xls', '.xlsx']:
            df = pd.read_excel(dataset_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Please provide CSV or Excel file.")
        
        filename = os.path.basename(dataset_path)

        # Try to load the model with retry logic
        try:
            model_name = "facebook/bart-large-mnli"
            model, tokenizer = load_model_with_retry(model_name, retries=3, timeout=1200)
            classifier = pipeline("zero-shot-classification", model=model, tokenizer=tokenizer, device=-1)
        except Exception as e:
            return jsonify({"error": f"Model loading failed: {str(e)}"}), 500

        candidate_labels = [
            "Financial data", "Retail data", "Stock data",
            "Inventory data", "Healthcare data", "Marketing data", "Sales data","College Data"
        ]

        # Classify based on dataset columns
        columns_text = " ".join(df.columns)
        rows_text = df.head().to_string(index=False)
        combined_text = columns_text + " " + rows_text
        model_prediction = classifier(combined_text, candidate_labels)

        if user_description:
            description_prediction = classifier(filename + user_description, candidate_labels)
            combined_scores = {
                label: (
                    0.6 * description_prediction["scores"][description_prediction["labels"].index(label)] +
                    0.4 * model_prediction["scores"][model_prediction["labels"].index(label)]
                )
                for label in candidate_labels
            }
            final_prediction = max(combined_scores, key=combined_scores.get)
        else:
            final_prediction = model_prediction["labels"][0]

        # Save result
        try:
            business_insight_path = os.path.join(output_dir, f"{user_id}_Business_Insight.txt")
            with open(business_insight_path, "w") as file:
                file.write(final_prediction)
            print(f"Business insight saved at {business_insight_path}")
        except Exception as e:
            print(f"Error writing business insight file: {str(e)}")
            raise

        # os.remove(dataset_path)
        return jsonify({
            "business_insight": final_prediction, 
            "output_path": business_insight_path
        }), 200
    except Exception as e:
        return jsonify({"error": f"Error during insight extraction: {str(e)}"}), 500


def detect_outliers(df):
    """Detect outliers in numeric columns and return stats + boxplot image"""
    outlier_info = {}

    for col in df.select_dtypes(include=[np.number]).columns:
        Q1 = df[col].quantile(0.25)
        Q3 = df[col].quantile(0.75)
        IQR = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        outliers = df[(df[col] < lower_bound) | (df[col] > upper_bound)][col]
            # Create boxplot image
        plt.figure(figsize=(8, 4))
        sns.boxplot(x=df[col])
        plt.title(f"Boxplot for {col}")
        plt.tight_layout()

        buffer = BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)
        plot_data = base64.b64encode(buffer.getvalue()).decode('utf-8')
        plt.close()

        outlier_info[col] = {
            'count': len(outliers),
            'plot': plot_data,
            'min': float(lower_bound),
            'max': float(upper_bound),
            'mean': float(df[col].mean()),
            'median': float(df[col].median())
        }

    return outlier_info



@app.route('/detect_outliers', methods=['POST'])
def detect_outliers_route():
    """Endpoint to detect outliers in the uploaded dataset"""
    try:
        # Use the latest dataset in the /data/ directory
        dataset_files = [f for f in os.listdir(data_dir) if f.endswith(('.csv', '.xlsx'))]
        if not dataset_files:
            return jsonify({"error": "No dataset found"}), 400

        dataset_path = os.path.join(data_dir, dataset_files[-1])
        _, ext = os.path.splitext(dataset_path)

        if ext == '.csv':
            df = pd.read_csv(dataset_path)
        elif ext in ['.xls', '.xlsx']:
            df = pd.read_excel(dataset_path)
        else:
            return jsonify({"error": "Unsupported file format"}), 400

        # Clean types and handle datetime before detecting outliers
        df = clean_column_types(df)
        df = handle_datetime_columns(df)

        # Detect and return outliers
        outlier_info = detect_outliers(df)
        return jsonify({"outliers": outlier_info}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500


from flask import request, jsonify
import pandas as pd
import os

@app.route('/handle_all_outliers', methods=['POST'])
def handle_all_outliers():
    try:
        data = request.get_json()
        method = data.get('method')  # 'remove', 'mean', 'median', 'keep'
        
        if not method:
            return jsonify({"error": "Missing 'method' in request"}), 400

        # Load latest mapped dataset or fallback
        dataset_path = os.path.join(data_dir, 'mapped_dataset.csv')
        if not os.path.exists(dataset_path):
            files = [f for f in os.listdir(data_dir) if f.endswith(('.csv', '.xlsx'))]
            if not files:
                return jsonify({'error': 'No dataset file found'}), 400
            dataset_path = os.path.join(data_dir, files[-1])

        _, ext = os.path.splitext(dataset_path)
        if ext == '.csv':
            df = pd.read_csv(dataset_path)
        elif ext in ['.xls', '.xlsx']:
            df = pd.read_excel(dataset_path)
        else:
            return jsonify({"error": "Unsupported file format"}), 400

        numeric_cols = df.select_dtypes(include='number').columns
        if numeric_cols.empty:
            return jsonify({"error": "No numeric columns found for outlier handling"}), 400

        for col in numeric_cols:
            Q1 = df[col].quantile(0.25)
            Q3 = df[col].quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR

            if method == 'Remove_outliers':
                df = df[(df[col] >= lower_bound) & (df[col] <= upper_bound)]
            elif method == 'Replace_with_mean':
                df[col] = df[col].mask((df[col] < lower_bound) | (df[col] > upper_bound), df[col].mean())
            elif method == 'Replace_with_median':
                df[col] = df[col].mask((df[col] < lower_bound) | (df[col] > upper_bound), df[col].median())
            elif method == 'Leave_as_is':
                continue
            else:
                return jsonify({"error": "Invalid method. Use 'remove', 'mean', 'median', or 'keep'."}), 400

        # Save the updated dataset
        df.to_csv(dataset_path, index=False)

        return jsonify({
            "status": "Outliers handled successfully for all numeric columns",
            "method": method,
            "columns_processed": list(numeric_cols)
        }), 200

    except Exception as e:
        return jsonify({"error": f"Error handling outliers: {str(e)}"}), 500


def handle_nulls_global_choice(df, choice):
    for col in df.columns:
        if df[col].isnull().sum() == 0:
            continue

        if choice == 'drop_column':
            df = df.drop(columns=[col])
        elif choice == 'drop_nulls':
            df = df.dropna(subset=[col])
        elif choice == 'fill_na' and df[col].dtype == 'object':
            df[col] = df[col].fillna('N/A')
        elif choice == 'fill_mean' and df[col].dtype != 'object':
            df[col] = df[col].fillna(df[col].mean())
        elif choice == 'fill_median' and df[col].dtype != 'object':
            df[col] = df[col].fillna(df[col].median())
        elif choice == 'fill_mode':
            df[col] = df[col].fillna(df[col].mode()[0])
    return df

@app.route('/handle_nulls', methods=['POST'])
def handle_nulls_route():
    """
    Endpoint to handle nulls using a global choice.
    Expects JSON input: {"choice": "fill_mean" | "drop_column" | ...}
    """
    try:
        data = request.get_json()
        choice = data.get("choice")

        if not choice:
            return jsonify({"error": "Missing 'choice' in request"}), 400

        # Load latest dataset
        dataset_files = [f for f in os.listdir(data_dir) if f.endswith(('.csv', '.xlsx'))]
        if not dataset_files:
            return jsonify({"error": "No dataset found"}), 400

        dataset_path = os.path.join(data_dir, dataset_files[-1])
        _, ext = os.path.splitext(dataset_path)

        if ext == '.csv':
            df = pd.read_csv(dataset_path)
        elif ext in ['.xls', '.xlsx']:
            df = pd.read_excel(dataset_path)
        else:
            return jsonify({"error": "Unsupported file format"}), 400

        # Clean column types and datetime formats
        df = clean_column_types(df)
        df = handle_datetime_columns(df)

        # Handle nulls using the global choice
        df = handle_nulls_global_choice(df, choice)

        # Save back the cleaned file (overwrite or change filename if preferred)
        df.to_csv(dataset_path, index=False)

        return jsonify({"message": f"Null handling using '{choice}' applied successfully."}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/map_columns', methods=['POST'])
def map_columns():
    business_column_dict = {
        "Financial data": ["transaction_id", "account_number", "balance", "transaction_amount", "interest_rate", "loan_amount", "payment_due", "credit_score", "investment", "roi", "debt", "mortgage", "savings", "tax", "budget", "income", "expense", "insurance", "liabilities", "assets", "net_worth", "cash_flow", "financial_goal", "dividends", "retirement_fund"],
        "Retail data": ["product_id", "product_name", "category", "price", "discount", "stock", "supplier", "revenue", "profit", "sales_volume", "customer_id", "order_id", "return_rate", "cart_abandonment", "promotion", "loyalty_points", "shelf_life", "store_location", "inventory_turnover", "market_trend", "advertising_spend", "online_vs_offline", "customer_feedback", "supplier_rating"],
        "Stock data": ["stock_symbol", "company_name", "market_cap", "opening_price", "closing_price", "volume", "52_week_high", "52_week_low", "pe_ratio", "eps", "dividend_yield", "shares_outstanding", "sector", "beta", "roe", "market_trend", "ipo_date", "analyst_rating", "institutional_ownership", "buy_sell_ratio", "news_sentiment", "stock_split", "insider_trading", "hedge_fund_activity"],
        "Inventory data": ["item_id", "item_name", "quantity", "reorder_level", "supplier", "warehouse_location", "cost_price", "selling_price", "storage_cost", "stock_rotation", "inventory_value", "backorder", "order_lead_time", "expiration_date", "stockout_rate", "safety_stock", "demand_forecast", "inventory_turnover", "obsolete_stock", "barcode", "sku", "batch_number", "quality_control", "inventory_audit"],
        "Healthcare data": ["patient_id", "age", "gender", "diagnosis", "treatment", "medication", "insurance", "hospital_id", "doctor_id", "medical_history", "allergies", "vital_signs", "admission_date", "discharge_date", "surgery", "lab_results", "follow_up", "emergency_contact", "billing_amount", "medical_tests", "chronic_disease", "hospital_rating", "vaccination_status", "patient_feedback", "readmission_rate"],
        "Marketing data": ["campaign_id", "ad_spend", "click_through_rate", "conversion_rate", "customer_acquisition_cost", "roi", "leads", "customer_lifetime_value", "email_open_rate", "social_media_engagement", "brand_awareness", "market_research", "competitor_analysis", "customer_feedback", "survey_results", "target_audience", "demographics", "geotargeting", "promotion", "loyalty_program", "influencer_partnership", "ad_impressions", "customer_retention", "pricing_strategy", "website_traffic"],
        "Sales data": ["sales_id", "customer_id", "product_id", "sales_amount", "profit_margin", "discount_applied", "region", "sales_channel", "order_date", "delivery_date", "customer_feedback", "refund", "customer_loyalty", "seasonal_trends", "market_share", "revenue_growth", "upselling", "cross_selling", "crm_data", "churn_rate", "repeat_customers", "average_order_value", "conversion_funnel", "customer_support", "purchase_frequency"],
        "College Data": ["student_id", "cgpa", "internships", "projects", "certifications", "aptitude_test_score", "soft_skills_rating", "extracurricular_activities", "placement_training", "ssc_marks", "hsc_marks", "placement_status", "major", "minor", "academic_year", "exam_score", "club_participation", "scholarship_status", "sports_activity", "volunteer_work", "faculty_rating", "assignment_score", "peer_review", "library_usage", "mentor_assigned"]
    }

    data = request.get_json()
    context = data.get("context")

    if not context:
        return jsonify({'error': 'Missing business context'}), 400

    if context not in business_column_dict:
        return jsonify({'error': 'Invalid business context'}), 400

    # Get the most recent file in the folder
    valid_files = [f for f in os.listdir(data_dir) if f.endswith(('.csv', '.xlsx', '.xls'))]
    if not valid_files:
        return jsonify({'error': 'No data files found'}), 400

    latest_file = max(valid_files, key=lambda f: os.path.getmtime(os.path.join(data_dir, f)))
    dataset_path = os.path.join(data_dir, latest_file)
    _, ext = os.path.splitext(dataset_path)

    try:
        if ext == '.csv':
            df = pd.read_csv(dataset_path)
        elif ext in ['.xls', '.xlsx']:
            df = pd.read_excel(dataset_path)
        else:
            return jsonify({"error": "Unsupported file format"}), 400
    except Exception as e:
        return jsonify({'error': f'Failed to read file: {str(e)}'}), 500

    available_columns = df.columns.tolist()
    expected_columns = business_column_dict[context]

    return jsonify({
        'file_name': latest_file,
        'available_columns': available_columns,
        'expected_columns': expected_columns
    })
@app.route('/submit_mapping', methods=['POST'])
def submit_mapping():
    data = request.get_json()
    mapping = data.get('mapping', {})

    # Find all dataset files in data_dir
    files = [f for f in os.listdir(data_dir) if f.endswith(('.csv', '.xlsx'))]
    if not files:
        return jsonify({'error': 'No dataset file found'}), 400

    # Load the latest dataset file
    dataset_path = os.path.join(data_dir, files[-1])
    _, ext = os.path.splitext(dataset_path)

    if ext == '.csv':
        df = pd.read_csv(dataset_path)
    elif ext in ['.xls', '.xlsx']:
        df = pd.read_excel(dataset_path)
    else:
        return jsonify({"error": "Unsupported file format"}), 400

    # Apply column mapping
    df.rename(columns=mapping, inplace=True)

    # Delete all existing dataset files in the directory
    # for f in os.listdir(data_dir):
    #     if f.endswith(('.csv', '.xlsx')):
    #         os.remove(os.path.join(data_dir, f))

    # Save the newly mapped dataset
    new_file_path = os.path.join(mapped_dir, 'mapped_dataset.csv')
    df.to_csv(new_file_path, index=False)

    return jsonify({
        "status": "Mapping applied, old files removed, and new dataset saved",
        "new_columns": df.columns.tolist()
    })





def save_plot_to_word(doc, fig, title):
    image_stream = BytesIO()
    fig.tight_layout()
    fig.savefig(image_stream, format='png')
    image_stream.seek(0)
    doc.add_paragraph(title)
    doc.add_picture(image_stream)
    image_stream.close()
    plt.close(fig)

def generate_financial_report(df, column_mapping, inv_column_mapping, doc):
    report = {}
    doc.add_heading('Financial Data Report', level=1)
    matched_columns_count = 0
    
    if 'balance' in column_mapping.values():
        report['Total Balance'] = df[inv_column_mapping['balance']].sum()
        doc.add_paragraph(f'Total Balance: {report["Total Balance"]}')
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.histplot(df[inv_column_mapping['balance']], bins=30, kde=True, ax=ax)
        ax.set_title('Balance Distribution')
        save_plot_to_word(doc, fig, 'Balance Distribution')
        matched_columns_count+=1
    
    if 'transaction_amount' in column_mapping.values():
        report['Total Transactions'] = df[inv_column_mapping['transaction_amount']].sum()
        doc.add_paragraph(f'Total Transactions: {report["Total Transactions"]}')
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.boxplot(x=df[inv_column_mapping['transaction_amount']], ax=ax)
        ax.set_title('Transaction Amount Distribution')
        save_plot_to_word(doc, fig, 'Transaction Amount Distribution')
        matched_columns_count+=1
    
    if 'credit_score' in column_mapping.values():
        report['Average Credit Score'] = df[inv_column_mapping['credit_score']].mean()
        doc.add_paragraph(f'Average Credit Score: {report["Average Credit Score"]}')
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.histplot(df[inv_column_mapping['credit_score']], bins=30, kde=True, ax=ax)
        ax.set_title('Credit Score Distribution')
        save_plot_to_word(doc, fig, 'Credit Score Distribution')
        matched_columns_count+=1
    
    if 'loan_amount' in column_mapping.values():
        report['Total Loan Amount'] = df[inv_column_mapping['loan_amount']].sum()
        doc.add_paragraph(f'Total Loan Amount: {report["Total Loan Amount"]}')
        matched_columns_count+=1
    
    if 'interest_rate' in column_mapping.values():
        report['Average Interest Rate'] = df[inv_column_mapping['interest_rate']].mean()
        doc.add_paragraph(f'Average Interest Rate: {report["Average Interest Rate"]}')
        matched_columns_count+=1
    
    if 'savings' in column_mapping.values():
        report['Total Savings'] = df[inv_column_mapping['savings']].sum()
        doc.add_paragraph(f'Total Savings: {report["Total Savings"]}')
        matched_columns_count+=1
    
    if 'investment' in column_mapping.values():
        report['Total Investment'] = df[inv_column_mapping['investment']].sum()
        doc.add_paragraph(f'Total Investment: {report["Total Investment"]}')
        matched_columns_count+=1
    
    if 'assets' in column_mapping.values() and 'liabilities' in column_mapping.values():
        report['Net Worth'] = df[inv_column_mapping['assets']].sum() - df[inv_column_mapping['liabilities']].sum()
        doc.add_paragraph(f'Net Worth: {report["Net Worth"]}')
        matched_columns_count+=1
    
    if 'income' in column_mapping.values():
        report['Total Income'] = df[inv_column_mapping['income']].sum()
        doc.add_paragraph(f'Total Income: {report["Total Income"]}')
        matched_columns_count+=1
    
    if 'expense' in column_mapping.values():
        report['Total Expenses'] = df[inv_column_mapping['expense']].sum()
        doc.add_paragraph(f'Total Expenses: {report["Total Expenses"]}')
        matched_columns_count+=1
    
    if 'mortgage' in column_mapping.values():
        report['Total Mortgage'] = df[inv_column_mapping['mortgage']].sum()
        doc.add_paragraph(f'Total Mortgage: {report["Total Mortgage"]}')
        matched_columns_count+=1
    
    if 'tax' in column_mapping.values():
        report['Total Tax'] = df[inv_column_mapping['tax']].sum()
        doc.add_paragraph(f'Total Tax: {report["Total Tax"]}')
        matched_columns_count+=1
    
    if 'insurance' in column_mapping.values():
        report['Total Insurance'] = df[inv_column_mapping['insurance']].sum()
        doc.add_paragraph(f'Total Insurance: {report["Total Insurance"]}')
        matched_columns_count+=1
    
    if 'cash_flow' in column_mapping.values():
        report['Cash Flow'] = df[inv_column_mapping['income']].sum() - df[inv_column_mapping['expense']].sum()
        doc.add_paragraph(f'Cash Flow: {report["Cash Flow"]}')
        matched_columns_count+=1
    
    if 'retirement_fund' in column_mapping.values():
        report['Total Retirement Fund'] = df[inv_column_mapping['retirement_fund']].sum()
        doc.add_paragraph(f'Total Retirement Fund: {report["Total Retirement Fund"]}')
        matched_columns_count+=1
    
    if matched_columns_count == 0:
        generate_plots_for_unmatched_columns(df, doc)
        
    return report

def generate_healthcare_report(df, column_mapping, hc_column_mapping, doc):
    report = {}
    doc.add_heading('Healthcare Data Report', level=1)
    matched_columns_count = 0

    if 'age' in column_mapping.values():
        report['Average Age'] = df[hc_column_mapping['age']].mean()
        doc.add_paragraph(f'Average Age of Patients: {report["Average Age"]:.2f}')
        matched_columns_count+=1

    if 'gender' in column_mapping.values():
        gender_counts = df[hc_column_mapping['gender']].value_counts()
        doc.add_paragraph('Gender Distribution:')
        for gender, count in gender_counts.items():
            doc.add_paragraph(f'{gender}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=gender_counts.index, y=gender_counts.values, ax=ax)
        ax.set_title('Gender Distribution')
        save_plot_to_word(doc, fig, 'Gender Distribution')
        matched_columns_count+=1

    if 'diagnosis' in column_mapping.values():
        diagnosis_counts = df[hc_column_mapping['diagnosis']].value_counts().head(10)
        doc.add_paragraph('Top 10 Diagnoses:')
        for diagnosis, count in diagnosis_counts.items():
            doc.add_paragraph(f'{diagnosis}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=diagnosis_counts.values, y=diagnosis_counts.index, ax=ax)
        ax.set_title('Top Diagnoses')
        save_plot_to_word(doc, fig, 'Top Diagnoses')
        matched_columns_count+=1

    if 'treatment' in column_mapping.values():
        treatment_counts = df[hc_column_mapping['treatment']].value_counts().head(10)
        doc.add_paragraph('Most Common Treatments:')
        for treatment, count in treatment_counts.items():
            doc.add_paragraph(f'{treatment}: {count}')
        matched_columns_count+=1

    if 'insurance' in column_mapping.values():
        matched_columns_count+=1
        insurance_counts = df[hc_column_mapping['insurance']].value_counts()
        doc.add_paragraph('Insurance Provider Distribution:')
        for provider, count in insurance_counts.items():
            doc.add_paragraph(f'{provider}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=insurance_counts.index, y=insurance_counts.values, ax=ax)
        ax.set_title('Insurance Providers')
        ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right')
        save_plot_to_word(doc, fig, 'Insurance Providers')

    if 'hospital_id' in column_mapping.values():
        matched_columns_count+=1
        hospital_counts = df[hc_column_mapping['hospital_id']].value_counts().head(10)
        doc.add_paragraph('Top Hospitals by Patient Count:')
        for hospital, count in hospital_counts.items():
            doc.add_paragraph(f'{hospital}: {count}')

    if 'billing_amount' in column_mapping.values():
        matched_columns_count+=1
        report['Average Billing Amount'] = df[hc_column_mapping['billing_amount']].mean()
        doc.add_paragraph(f'Average Billing Amount: {report["Average Billing Amount"]:.2f}')
        fig, ax = plt.subplots()
        sns.histplot(df[hc_column_mapping['billing_amount']], bins=30, kde=True, ax=ax)
        ax.set_title('Billing Amount Distribution')
        save_plot_to_word(doc, fig, 'Billing Amount Distribution')

    if 'lab_results' in column_mapping.values():
        matched_columns_count+=1
        lab_result_nulls = df[hc_column_mapping['lab_results']].isnull().sum()
        doc.add_paragraph(f'Number of Missing Lab Results: {lab_result_nulls}')

    if 'readmission_rate' in column_mapping.values():
        matched_columns_count+=1
        report['Average Readmission Rate'] = df[hc_column_mapping['readmission_rate']].mean()
        doc.add_paragraph(f'Average Readmission Rate: {report["Average Readmission Rate"]:.2f}')

    if 'vaccination_status' in column_mapping.values():
        matched_columns_count+=1      
        vaccination_counts = df[hc_column_mapping['vaccination_status']].value_counts()
        doc.add_paragraph('Vaccination Status Distribution:')
        for status, count in vaccination_counts.items():
            doc.add_paragraph(f'{status}: {count}')

    if 'patient_feedback' in column_mapping.values():
        matched_columns_count+=1
        feedback_counts = df[hc_column_mapping['patient_feedback']].value_counts().head(5)
        doc.add_paragraph('Top Patient Feedback:')
        for feedback, count in feedback_counts.items():
            doc.add_paragraph(f'{feedback}: {count}')

    if matched_columns_count == 0:
        generate_plots_for_unmatched_columns(df, doc)
        
    return report

def generate_inventory_report(df, column_mapping, inv_column_mapping, doc):
    matched_columns_count = 0
    report = {}
    doc.add_heading('Inventory Data Report', level=1)

    if 'quantity' in column_mapping.values():
        matched_columns_count+=1
        report['Total Quantity'] = df[inv_column_mapping['quantity']].sum()
        doc.add_paragraph(f'Total Quantity: {report["Total Quantity"]}')

    if 'inventory_value' in column_mapping.values():
        matched_columns_count+=1
        report['Total Inventory Value'] = df[inv_column_mapping['inventory_value']].sum()
        doc.add_paragraph(f'Total Inventory Value: {report["Total Inventory Value"]}')

    if 'cost_price' in column_mapping.values() and 'selling_price' in column_mapping.values():
        matched_columns_count+=1
        df['profit_margin'] = df[inv_column_mapping['selling_price']] - df[inv_column_mapping['cost_price']]
        report['Average Profit Margin'] = df['profit_margin'].mean()
        doc.add_paragraph(f'Average Profit Margin: {report["Average Profit Margin"]:.2f}')
        fig, ax = plt.subplots()
        sns.histplot(df['profit_margin'], bins=30, kde=True, ax=ax)
        ax.set_title('Profit Margin Distribution')
        save_plot_to_word(doc, fig, 'Profit Margin Distribution')

    if 'supplier' in column_mapping.values():
        matched_columns_count+=1
        supplier_counts = df[inv_column_mapping['supplier']].value_counts()
        doc.add_paragraph('Supplier Distribution:')
        for supplier, count in supplier_counts.items():
            doc.add_paragraph(f'{supplier}: {count}')
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.barplot(x=supplier_counts.index, y=supplier_counts.values, ax=ax)
        ax.set_title('Supplier Distribution')
        ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right')
        save_plot_to_word(doc, fig, 'Supplier Distribution')

    if 'reorder_level' in column_mapping.values():
        matched_columns_count+=1
        below_reorder = df[df[inv_column_mapping['quantity']] < df[inv_column_mapping['reorder_level']]]
        report['Items Below Reorder Level'] = below_reorder.shape[0]
        doc.add_paragraph(f'Items Below Reorder Level: {report["Items Below Reorder Level"]}')

    if 'stockout_rate' in column_mapping.values():
        matched_columns_count+=1
        report['Average Stockout Rate'] = df[inv_column_mapping['stockout_rate']].mean()
        doc.add_paragraph(f'Average Stockout Rate: {report["Average Stockout Rate"]:.2f}')

    if 'inventory_turnover' in column_mapping.values():
        matched_columns_count+=1
        report['Average Inventory Turnover'] = df[inv_column_mapping['inventory_turnover']].mean()
        doc.add_paragraph(f'Average Inventory Turnover: {report["Average Inventory Turnover"]:.2f}')
        fig, ax = plt.subplots()
        sns.histplot(df[inv_column_mapping['inventory_turnover']], bins=30, kde=True, ax=ax)
        ax.set_title('Inventory Turnover Distribution')
        save_plot_to_word(doc, fig, 'Inventory Turnover Distribution')

    if 'demand_forecast' in column_mapping.values():
        matched_columns_count+=1
        report['Average Demand Forecast'] = df[inv_column_mapping['demand_forecast']].mean()
        doc.add_paragraph(f'Average Demand Forecast: {report["Average Demand Forecast"]:.2f}')

    if 'obsolete_stock' in column_mapping.values():
        matched_columns_count+=1
        report['Obsolete Stock Count'] = df[df[inv_column_mapping['obsolete_stock']] > 0].shape[0]
        doc.add_paragraph(f'Obsolete Stock Count: {report["Obsolete Stock Count"]}')

    if 'warehouse_location' in column_mapping.values():
        matched_columns_count+=1
        location_counts = df[inv_column_mapping['warehouse_location']].value_counts()
        doc.add_paragraph('Warehouse Location Distribution:')
        for location, count in location_counts.items():
            doc.add_paragraph(f'{location}: {count}')
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.barplot(x=location_counts.index, y=location_counts.values, ax=ax)
        ax.set_title('Warehouse Locations')
        ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right')
        save_plot_to_word(doc, fig, 'Warehouse Location Distribution')

    if 'expiration_date' in column_mapping.values():
        matched_columns_count+=1
        df['expiration_date_parsed'] = pd.to_datetime(df[inv_column_mapping['expiration_date']], errors='coerce')
        expiring_soon = df[df['expiration_date_parsed'] < pd.Timestamp.today() + pd.Timedelta(days=30)]
        report['Items Expiring Soon'] = expiring_soon.shape[0]
        doc.add_paragraph(f'Items Expiring Within 30 Days: {report["Items Expiring Soon"]}')

    if matched_columns_count == 0:
        generate_plots_for_unmatched_columns(df, doc)
        
    return report

def generate_retail_report(df, column_mapping, inv_column_mapping, doc):
    matched_columns_count = 0
    report = {}
    doc.add_heading('Retail Data Report', level=1)
    
    if 'revenue' in column_mapping.values():
        matched_columns_count+=1
        report['Total Revenue'] = df[inv_column_mapping['revenue']].sum()
        doc.add_paragraph(f'Total Revenue: {report["Total Revenue"]}')
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.histplot(df[inv_column_mapping['revenue']], bins=30, kde=True, ax=ax)
        ax.set_title('Revenue Distribution')
        save_plot_to_word(doc, fig, 'Revenue Distribution')
    
    if 'profit' in column_mapping.values():
        matched_columns_count+=1
        report['Total Profit'] = df[inv_column_mapping['profit']].sum()
        doc.add_paragraph(f'Total Profit: {report["Total Profit"]}')
    
    if 'sales_volume' in column_mapping.values():
        matched_columns_count+=1
        report['Total Sales Volume'] = df[inv_column_mapping['sales_volume']].sum()
        doc.add_paragraph(f'Total Sales Volume: {report["Total Sales Volume"]}')
        
        # Top 10 best-selling products
        if 'product_name' in column_mapping.values():
            top_products = df.groupby(inv_column_mapping['product_name'])[inv_column_mapping['sales_volume']].sum().nlargest(10)
            fig, ax = plt.subplots(figsize=(10, 5))
            top_products.plot(kind='bar', ax=ax)
            ax.set_title('Top 10 Best-Selling Products')
            save_plot_to_word(doc, fig, 'Top 10 Best-Selling Products')
    
    if 'stock' in column_mapping.values():
        matched_columns_count+=1
        report['Average Stock Level'] = df[inv_column_mapping['stock']].mean()
        doc.add_paragraph(f'Average Stock Level: {report["Average Stock Level"]}')
        
        # Stock Level Distribution
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.histplot(df[inv_column_mapping['stock']], bins=30, kde=True, ax=ax)
        ax.set_title('Stock Level Distribution')
        save_plot_to_word(doc, fig, 'Stock Level Distribution')
    
    if 'return_rate' in column_mapping.values():
        matched_columns_count+=1
        report['Average Return Rate'] = df[inv_column_mapping['return_rate']].mean()
        doc.add_paragraph(f'Average Return Rate: {report["Average Return Rate"]}')
        
        # Return Rate vs Profit Analysis
        if 'profit' in column_mapping.values():
            fig, ax = plt.subplots(figsize=(10, 5))
            sns.boxplot(x=df[inv_column_mapping['return_rate']], y=df[inv_column_mapping['profit']], ax=ax)
            ax.set_title('Return Rate Impact on Profit')
            save_plot_to_word(doc, fig, 'Return Rate Impact on Profit')
    
    if 'advertising_spend' in column_mapping.values() and 'revenue' in column_mapping.values():
        matched_columns_count+=1
        report['Total Advertising Spend'] = df[inv_column_mapping['advertising_spend']].sum()
        doc.add_paragraph(f'Total Advertising Spend: {report["Total Advertising Spend"]}')
        
        # Advertising Spend vs Revenue
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.regplot(x=df[inv_column_mapping['advertising_spend']], y=df[inv_column_mapping['revenue']], ax=ax)
        ax.set_title('Advertising Spend vs Revenue')
        save_plot_to_word(doc, fig, 'Advertising Spend vs Revenue')
    
    if matched_columns_count == 0:
        generate_plots_for_unmatched_columns(df, doc)
        
    return report

def generate_college_report(df, column_mapping, inv_column_mapping, doc):
    matched_columns_count = 0
    report = {}
    doc.add_heading('College Data Report', level=1)
    
    if 'cgpa' in column_mapping.values():
        matched_columns_count+=1
        report['Average CGPA'] = df[inv_column_mapping['cgpa']].mean()
        doc.add_paragraph(f'Average CGPA: {report["Average CGPA"]}')
        fig, ax = plt.subplots()
        sns.histplot(df[inv_column_mapping['cgpa']], bins=30, kde=True, ax=ax)
        ax.set_title('CGPA Distribution')
        save_plot_to_word(doc, fig, 'CGPA Distribution')
    
    if 'internships' in column_mapping.values():
        matched_columns_count+=1
        report['Average Internships Completed'] = df[inv_column_mapping['internships']].mean()
        doc.add_paragraph(f'Average Internships Completed: {report["Average Internships Completed"]}')
    
    if 'placement_status' in column_mapping.values():
        matched_columns_count+=1
        placement_counts = df[inv_column_mapping['placement_status']].value_counts()
        doc.add_paragraph('Placement Status Distribution:')
        for status, count in placement_counts.items():
            doc.add_paragraph(f'{status}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=placement_counts.index, y=placement_counts.values, ax=ax)
        ax.set_title('Placement Status Distribution')
        save_plot_to_word(doc, fig, 'Placement Status Distribution')
    
    if 'exam_score' in column_mapping.values():
        matched_columns_count+=1
        report['Average Exam Score'] = df[inv_column_mapping['exam_score']].mean()
        doc.add_paragraph(f'Average Exam Score: {report["Average Exam Score"]}')
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.histplot(df[inv_column_mapping['exam_score']], bins=30, kde=True, ax=ax)
        ax.set_title('Exam Score Distribution')
        save_plot_to_word(doc, fig, 'Exam Score Distribution')
    
    if 'extracurricular_activities' in column_mapping.values():
        matched_columns_count+=1
        report['Participation in Extracurricular Activities'] = df[inv_column_mapping['extracurricular_activities']].notnull().sum()
        doc.add_paragraph(f'Total Extracurricular Activities Participation: {report["Participation in Extracurricular Activities"]}')
    
    if 'faculty_rating' in column_mapping.values():
        matched_columns_count+=1
        report['Average Faculty Rating'] = df[inv_column_mapping['faculty_rating']].mean()
        doc.add_paragraph(f'Average Faculty Rating: {report["Average Faculty Rating"]}')
    
    if 'peer_review' in column_mapping.values():
        matched_columns_count+=1
        report['Average Peer Review Score'] = df[inv_column_mapping['peer_review']].mean()
        doc.add_paragraph(f'Average Peer Review Score: {report["Average Peer Review Score"]}')
    
    if 'library_usage' in column_mapping.values():
        matched_columns_count+=1
        report['Average Library Usage'] = df[inv_column_mapping['library_usage']].mean()
        doc.add_paragraph(f'Average Library Usage: {report["Average Library Usage"]}')
    
    if 'ssc_marks' in column_mapping.values():
        matched_columns_count+=1
        report['Average SSC Marks'] = df[inv_column_mapping['ssc_marks']].mean()
        doc.add_paragraph(f'Average SSC Marks: {report["Average SSC Marks"]}')
    
    if 'hsc_marks' in column_mapping.values():
        matched_columns_count+=1
        report['Average HSC Marks'] = df[inv_column_mapping['hsc_marks']].mean()
        doc.add_paragraph(f'Average HSC Marks: {report["Average HSC Marks"]}')
    
    if 'scholarship_status' in column_mapping.values():
        matched_columns_count+=1
        scholarship_counts = df[inv_column_mapping['scholarship_status']].value_counts()
        doc.add_paragraph('Scholarship Status Distribution:')
        for status, count in scholarship_counts.items():
            doc.add_paragraph(f'{status}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=scholarship_counts.index, y=scholarship_counts.values, ax=ax)
        ax.set_title('Scholarship Status Distribution')
        save_plot_to_word(doc, fig, 'Scholarship Status Distribution')
    
    if 'mentor_assigned' in column_mapping.values():
        matched_columns_count+=1
        mentor_counts = df[inv_column_mapping['mentor_assigned']].value_counts()
        doc.add_paragraph('Mentor Assignment Distribution:')
        for status, count in mentor_counts.items():
            doc.add_paragraph(f'{status}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=mentor_counts.index, y=mentor_counts.values, ax=ax)
        ax.set_title('Mentor Assignment Distribution')
        save_plot_to_word(doc, fig, 'Mentor Assignment Distribution')
    
    if 'placement_training' in column_mapping.values():
        matched_columns_count+=1
        report['Students Enrolled in Placement Training'] = df[inv_column_mapping['placement_training']].notnull().sum()
        doc.add_paragraph(f'Students Enrolled in Placement Training: {report["Students Enrolled in Placement Training"]}')
    
    if 'soft_skills_rating' in column_mapping.values():
        matched_columns_count+=1
        report['Average Soft Skills Rating'] = df[inv_column_mapping['soft_skills_rating']].mean()
        doc.add_paragraph(f'Average Soft Skills Rating: {report["Average Soft Skills Rating"]}')
    
    if 'volunteer_work' in column_mapping.values():
        matched_columns_count+=1
        report['Students Involved in Volunteer Work'] = df[inv_column_mapping['volunteer_work']].notnull().sum()
        doc.add_paragraph(f'Students Involved in Volunteer Work: {report["Students Involved in Volunteer Work"]}')
    
    if matched_columns_count == 0:
        generate_plots_for_unmatched_columns(df, doc)
        
    print(report)
    return report

def generate_stock_report(df, column_mapping, inv_column_mapping, doc):
    matched_columns_count = 0
    report = {}
    doc.add_heading('Stock Data Report', level=1)

    if 'market_cap' in column_mapping.values():
        matched_columns_count+=1
        report['Average Market Cap'] = df[inv_column_mapping['market_cap']].mean()
        doc.add_paragraph(f'Average Market Cap: {report["Average Market Cap"]}')

    if 'opening_price' in column_mapping.values() and 'closing_price' in column_mapping.values():
        matched_columns_count+=1
        df['daily_change'] = df[inv_column_mapping['closing_price']] - df[inv_column_mapping['opening_price']]
        report['Average Daily Change'] = df['daily_change'].mean()
        doc.add_paragraph(f'Average Daily Price Change: {report["Average Daily Change"]}')

    if 'volume' in column_mapping.values():
        matched_columns_count+=1
        report['Average Volume'] = df[inv_column_mapping['volume']].mean()
        doc.add_paragraph(f'Average Volume: {report["Average Volume"]}')

    if 'pe_ratio' in column_mapping.values():
        matched_columns_count+=1
        report['Average P/E Ratio'] = df[inv_column_mapping['pe_ratio']].mean()
        doc.add_paragraph(f'Average P/E Ratio: {report["Average P/E Ratio"]}')

    if 'eps' in column_mapping.values():
        matched_columns_count+=1
        report['Average EPS'] = df[inv_column_mapping['eps']].mean()
        doc.add_paragraph(f'Average EPS: {report["Average EPS"]}')

    if 'dividend_yield' in column_mapping.values():
        matched_columns_count+=1
        report['Average Dividend Yield'] = df[inv_column_mapping['dividend_yield']].mean()
        doc.add_paragraph(f'Average Dividend Yield: {report["Average Dividend Yield"]}')

    if 'sector' in column_mapping.values():
        matched_columns_count+=1
        sector_counts = df[inv_column_mapping['sector']].value_counts()
        doc.add_paragraph('Sector Distribution:')
        for sector, count in sector_counts.items():
            doc.add_paragraph(f'{sector}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=sector_counts.index, y=sector_counts.values, ax=ax)
        ax.set_title('Sector Distribution')
        ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right')
        save_plot_to_word(doc, fig, 'Sector Distribution')

    if 'beta' in column_mapping.values():
        matched_columns_count+=1
        report['Average Beta'] = df[inv_column_mapping['beta']].mean()
        doc.add_paragraph(f'Average Beta: {report["Average Beta"]}')

    if 'roe' in column_mapping.values():
        matched_columns_count+=1
        report['Average ROE'] = df[inv_column_mapping['roe']].mean()
        doc.add_paragraph(f'Average ROE: {report["Average ROE"]}')

    if 'news_sentiment' in column_mapping.values():
        matched_columns_count+=1
        sentiment_counts = df[inv_column_mapping['news_sentiment']].value_counts()
        doc.add_paragraph('News Sentiment Distribution:')
        for sentiment, count in sentiment_counts.items():
            doc.add_paragraph(f'{sentiment}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=sentiment_counts.index, y=sentiment_counts.values, ax=ax)
        ax.set_title('News Sentiment Distribution')
        save_plot_to_word(doc, fig, 'News Sentiment Distribution')

    if 'analyst_rating' in column_mapping.values():
        matched_columns_count+=1
        rating_counts = df[inv_column_mapping['analyst_rating']].value_counts()
        doc.add_paragraph('Analyst Rating Distribution:')
        for rating, count in rating_counts.items():
            doc.add_paragraph(f'{rating}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=rating_counts.index, y=rating_counts.values, ax=ax)
        ax.set_title('Analyst Rating Distribution')
        save_plot_to_word(doc, fig, 'Analyst Rating Distribution')

    if matched_columns_count == 0:
        generate_plots_for_unmatched_columns(df, doc)
        
    return report

def generate_marketing_report(df, column_mapping, mk_column_mapping, doc):
    matched_columns_count = 0
    report = {}
    doc.add_heading('Marketing Data Report', level=1)

    if 'ad_spend' in column_mapping.values():
        matched_columns_count+=1
        total_ad_spend = df[mk_column_mapping['ad_spend']].sum()
        avg_ad_spend = df[mk_column_mapping['ad_spend']].mean()
        report['Total Ad Spend'] = total_ad_spend
        report['Average Ad Spend'] = avg_ad_spend
        doc.add_paragraph(f"Total Ad Spend: ${total_ad_spend:,.2f}")
        doc.add_paragraph(f"Average Ad Spend: ${avg_ad_spend:,.2f}")

    if 'click_through_rate' in column_mapping.values():
        matched_columns_count+=1
        avg_ctr = df[mk_column_mapping['click_through_rate']].mean()
        report['Average Click Through Rate'] = avg_ctr
        doc.add_paragraph(f"Average Click Through Rate (CTR): {avg_ctr:.2%}")

    if 'conversion_rate' in column_mapping.values():
        matched_columns_count+=1
        avg_conversion = df[mk_column_mapping['conversion_rate']].mean()
        report['Average Conversion Rate'] = avg_conversion
        doc.add_paragraph(f"Average Conversion Rate: {avg_conversion:.2%}")

    if 'customer_acquisition_cost' in column_mapping.values():
        matched_columns_count+=1
        avg_cac = df[mk_column_mapping['customer_acquisition_cost']].mean()
        report['Average CAC'] = avg_cac
        doc.add_paragraph(f"Average Customer Acquisition Cost (CAC): ${avg_cac:,.2f}")

    if 'roi' in column_mapping.values():
        matched_columns_count+=1
        avg_roi = df[mk_column_mapping['roi']].mean()
        report['Average ROI'] = avg_roi
        doc.add_paragraph(f"Average Return on Investment (ROI): {avg_roi:.2f}")

    if 'customer_lifetime_value' in column_mapping.values():
        matched_columns_count+=1
        avg_clv = df[mk_column_mapping['customer_lifetime_value']].mean()
        report['Average CLV'] = avg_clv
        doc.add_paragraph(f"Average Customer Lifetime Value (CLV): ${avg_clv:,.2f}")

    if 'email_open_rate' in column_mapping.values():
        matched_columns_count+=1
        avg_open = df[mk_column_mapping['email_open_rate']].mean()
        report['Average Email Open Rate'] = avg_open
        doc.add_paragraph(f"Average Email Open Rate: {avg_open:.2%}")

    if 'social_media_engagement' in column_mapping.values():
        matched_columns_count+=1
        sm_counts = df[mk_column_mapping['social_media_engagement']].value_counts().head(5)
        doc.add_paragraph('Top Social Media Engagement Platforms:')
        for platform, count in sm_counts.items():
            doc.add_paragraph(f"{platform}: {count}")

    if 'brand_awareness' in column_mapping.values():
        matched_columns_count+=1
        avg_awareness = df[mk_column_mapping['brand_awareness']].mean()
        report['Average Brand Awareness'] = avg_awareness
        doc.add_paragraph(f"Average Brand Awareness Score: {avg_awareness:.2f}")

    if 'market_research' in column_mapping.values():
        matched_columns_count+=1
        insights = df[mk_column_mapping['market_research']].dropna().unique()[:5]
        doc.add_paragraph('Key Market Research Insights:')
        for insight in insights:
            doc.add_paragraph(f"- {insight}")

    if 'competitor_analysis' in column_mapping.values():
        matched_columns_count+=1
        competitors = df[mk_column_mapping['competitor_analysis']].dropna().unique()[:5]
        doc.add_paragraph('Top Competitor Mentions:')
        for competitor in competitors:
            doc.add_paragraph(f"- {competitor}")

    if 'website_traffic' in column_mapping.values():
        matched_columns_count+=1
        avg_traffic = df[mk_column_mapping['website_traffic']].mean()
        doc.add_paragraph(f"Average Website Traffic: {avg_traffic:,.0f} visits")

    if 'customer_feedback' in column_mapping.values():
        matched_columns_count+=1
        feedback_counts = df[mk_column_mapping['customer_feedback']].value_counts().head(5)
        doc.add_paragraph('Frequent Customer Feedback Themes:')
        for feedback, count in feedback_counts.items():
            doc.add_paragraph(f"{feedback}: {count}")

    if 'customer_retention' in column_mapping.values():
        matched_columns_count+=1
        avg_retention = df[mk_column_mapping['customer_retention']].mean()
        report['Customer Retention Rate'] = avg_retention
        doc.add_paragraph(f"Average Customer Retention Rate: {avg_retention:.2%}")

    if 'pricing_strategy' in column_mapping.values():
        matched_columns_count+=1
        pricing_strategies = df[mk_column_mapping['pricing_strategy']].value_counts().head(5)
        doc.add_paragraph('Top Pricing Strategies Adopted:')
        for strategy, count in pricing_strategies.items():
            doc.add_paragraph(f"{strategy}: {count}")

    if 'demographics' in column_mapping.values():
        matched_columns_count+=1
        demographics = df[mk_column_mapping['demographics']].dropna().unique()[:5]
        doc.add_paragraph('Key Demographic Segments:')
        for segment in demographics:
            doc.add_paragraph(f"- {segment}")

    if 'loyalty_program' in column_mapping.values():
        matched_columns_count+=1
        loyalty_counts = df[mk_column_mapping['loyalty_program']].value_counts().head(5)
        doc.add_paragraph('Loyalty Program Participation:')
        for program, count in loyalty_counts.items():
            doc.add_paragraph(f"{program}: {count}")

    if matched_columns_count == 0:
        generate_plots_for_unmatched_columns(df, doc)
        
    return report

def generate_sales_report(df, column_mapping, sales_column_mapping, doc):
    matched_columns_count = 0
    report = {}
    doc.add_heading('Sales Data Report', level=1)

    if 'sales_amount' in column_mapping.values():
        matched_columns_count+=1
        total_sales = df[sales_column_mapping['sales_amount']].sum()
        avg_sales = df[sales_column_mapping['sales_amount']].mean()
        report['Total Sales'] = total_sales
        report['Average Sales'] = avg_sales
        doc.add_paragraph(f"Total Sales Amount: ${total_sales:.2f}")
        doc.add_paragraph(f"Average Sales Amount: ${avg_sales:.2f}")

    if 'profit_margin' in column_mapping.values():
        matched_columns_count+=1
        avg_margin = df[sales_column_mapping['profit_margin']].mean()
        report['Average Profit Margin'] = avg_margin
        doc.add_paragraph(f"Average Profit Margin: {avg_margin:.2%}")

    if 'discount_applied' in column_mapping.values():
        matched_columns_count+=1
        discount_rate = df[sales_column_mapping['discount_applied']].mean()
        report['Average Discount Applied'] = discount_rate
        doc.add_paragraph(f"Average Discount Applied: {discount_rate:.2%}")

    if 'region' in column_mapping.values():
        matched_columns_count+=1
        doc.add_paragraph('Top Performing Regions:')
        region_sales = df.groupby(sales_column_mapping['region'])[sales_column_mapping['sales_amount']].sum().nlargest(5)
        for region, amount in region_sales.items():
            doc.add_paragraph(f"{region}: ${amount:.2f}")

    if 'sales_channel' in column_mapping.values():
        matched_columns_count+=1
        doc.add_paragraph('Sales by Channel:')
        channel_sales = df[sales_column_mapping['sales_channel']].value_counts()
        for channel, count in channel_sales.items():
            doc.add_paragraph(f"{channel}: {count}")

    if 'order_date' in column_mapping.values():
        matched_columns_count+=1
        df[sales_column_mapping['order_date']] = pd.to_datetime(df[sales_column_mapping['order_date']])
        monthly_sales = df.groupby(df[sales_column_mapping['order_date']].dt.to_period('M'))[sales_column_mapping['sales_amount']].sum()
        fig, ax = plt.subplots()
        monthly_sales.plot(ax=ax, kind='line', marker='o', title='Monthly Sales Trend')
        save_plot_to_word(doc, fig, "Monthly Sales Trend")

    if 'customer_feedback' in column_mapping.values():
        matched_columns_count+=1
        feedback = df[sales_column_mapping['customer_feedback']].value_counts().head(5)
        doc.add_paragraph('Top Customer Feedback Themes:')
        for theme, count in feedback.items():
            doc.add_paragraph(f"{theme}: {count}")

    if 'churn_rate' in column_mapping.values():
        matched_columns_count+=1
        avg_churn = df[sales_column_mapping['churn_rate']].mean()
        report['Average Churn Rate'] = avg_churn
        doc.add_paragraph(f"Average Customer Churn Rate: {avg_churn:.2%}")

    if 'repeat_customers' in column_mapping.values():
        matched_columns_count+=1
        repeat_rate = df[sales_column_mapping['repeat_customers']].mean()
        report['Repeat Customer Rate'] = repeat_rate
        doc.add_paragraph(f"Repeat Customer Rate: {repeat_rate:.2%}")

    if 'average_order_value' in column_mapping.values():
        matched_columns_count+=1
        avg_order_value = df[sales_column_mapping['average_order_value']].mean()
        report['Average Order Value'] = avg_order_value
        doc.add_paragraph(f"Average Order Value: ${avg_order_value:.2f}")

    if 'crm_data' in column_mapping.values():
        matched_columns_count+=1
        doc.add_paragraph('CRM Notes Samples:')
        notes = df[sales_column_mapping['crm_data']].dropna().unique()[:5]
        for note in notes:
            doc.add_paragraph(f"- {note}")

    if 'conversion_funnel' in column_mapping.values():
        matched_columns_count+=1
        funnel_stages = df[sales_column_mapping['conversion_funnel']].value_counts()
        doc.add_paragraph('Conversion Funnel Breakdown:')
        for stage, count in funnel_stages.items():
            doc.add_paragraph(f"{stage}: {count}")

    if 'purchase_frequency' in column_mapping.values():
        matched_columns_count+=1
        freq = df[sales_column_mapping['purchase_frequency']].mean()
        report['Average Purchase Frequency'] = freq
        doc.add_paragraph(f"Average Purchase Frequency: {freq:.2f} purchases per customer")

    if matched_columns_count == 0:
        generate_plots_for_unmatched_columns(df, doc)

    return report

def generate_plots_for_unmatched_columns(df, doc):
    """
    This function generates plots for numerical and categorical columns
    when no specific columns are matched in the if conditions.
    It generates suitable plots and adds them to the DOCX document.
    """
    # For numerical columns, generate histograms unless they are treated as categorical
    num_columns = df.select_dtypes(include=[np.number]).columns
    if num_columns.any():
        for col in num_columns:
            # If the column has fewer than 10 unique values, treat it as categorical
            if df[col].nunique() < 10:
                # Treat as categorical and generate a pie chart or bar plot
                value_counts = df[col].value_counts()
                if len(value_counts) <= 10:  # For smaller unique sets, pie chart
                    fig, ax = plt.subplots(figsize=(8, 5))
                    ax.pie(value_counts, labels=value_counts.index, autopct='%1.1f%%', startangle=90)
                    ax.set_title(f'Distribution of {col}')
                    save_plot_to_word(doc, fig, f'Distribution of {col}')
                else:  # For more than 10 unique values, bar plot
                    fig, ax = plt.subplots(figsize=(8, 5))
                    sns.barplot(x=value_counts.index, y=value_counts.values, ax=ax)
                    ax.set_title(f'Top Categories of {col}')
                    ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right')
                    save_plot_to_word(doc, fig, f'Top Categories of {col}')
            else:
                # Histogram for numerical columns with continuous values
                fig, ax = plt.subplots(figsize=(8, 5))
                sns.histplot(df[col], kde=True, ax=ax)
                ax.set_title(f'Distribution of {col}')
                save_plot_to_word(doc, fig, f'Distribution of {col}')

    # For non-numerical columns, check if they are suitable for categorical representation
    non_num_columns = df.select_dtypes(exclude=[np.number]).columns
    if non_num_columns.any():
        for col in non_num_columns:
            # If the column has fewer than 20 unique values, treat it as categorical
            if df[col].nunique() < 20:
                # Pie chart for categorical columns with fewer unique values
                value_counts = df[col].value_counts()
                fig, ax = plt.subplots(figsize=(8, 5))
                ax.pie(value_counts, labels=value_counts.index, autopct='%1.1f%%', startangle=90)
                ax.set_title(f'Distribution of {col}')
                save_plot_to_word(doc, fig, f'Distribution of {col}')
            else:
                # For columns with more unique values, skip pie chart, but can use a bar chart
                value_counts = df[col].value_counts().head(10)  # Limit to top 10 for better readability
                fig, ax = plt.subplots(figsize=(8, 5))
                sns.barplot(x=value_counts.index, y=value_counts.values, ax=ax)
                ax.set_title(f'Top 10 {col} Categories')
                ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right')
                save_plot_to_word(doc, fig, f'Top 10 {col} Categories')

    # Optional: Return a message to indicate that plots have been added to the report
    return "Plots for unmatched columns have been generated and saved."
def generate_stock_report(df, column_mapping, inv_column_mapping, doc):
    matched_columns_count = 0
    report = {}
    doc.add_heading('Stock Data Report', level=1)

    if 'market_cap' in column_mapping.values():
        matched_columns_count+=1
        report['Average Market Cap'] = df[inv_column_mapping['market_cap']].mean()
        doc.add_paragraph(f'Average Market Cap: {report["Average Market Cap"]}')

    if 'opening_price' in column_mapping.values() and 'closing_price' in column_mapping.values():
        matched_columns_count+=1
        df['daily_change'] = df[inv_column_mapping['closing_price']] - df[inv_column_mapping['opening_price']]
        report['Average Daily Change'] = df['daily_change'].mean()
        doc.add_paragraph(f'Average Daily Price Change: {report["Average Daily Change"]}')

    if 'volume' in column_mapping.values():
        matched_columns_count+=1
        report['Average Volume'] = df[inv_column_mapping['volume']].mean()
        doc.add_paragraph(f'Average Volume: {report["Average Volume"]}')

    if 'pe_ratio' in column_mapping.values():
        matched_columns_count+=1
        report['Average P/E Ratio'] = df[inv_column_mapping['pe_ratio']].mean()
        doc.add_paragraph(f'Average P/E Ratio: {report["Average P/E Ratio"]}')

    if 'eps' in column_mapping.values():
        matched_columns_count+=1
        report['Average EPS'] = df[inv_column_mapping['eps']].mean()
        doc.add_paragraph(f'Average EPS: {report["Average EPS"]}')

    if 'dividend_yield' in column_mapping.values():
        matched_columns_count+=1
        report['Average Dividend Yield'] = df[inv_column_mapping['dividend_yield']].mean()
        doc.add_paragraph(f'Average Dividend Yield: {report["Average Dividend Yield"]}')

    if 'sector' in column_mapping.values():
        matched_columns_count+=1
        sector_counts = df[inv_column_mapping['sector']].value_counts()
        doc.add_paragraph('Sector Distribution:')
        for sector, count in sector_counts.items():
            doc.add_paragraph(f'{sector}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=sector_counts.index, y=sector_counts.values, ax=ax)
        ax.set_title('Sector Distribution')
        ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right')
        save_plot_to_word(doc, fig, 'Sector Distribution')

    if 'beta' in column_mapping.values():
        matched_columns_count+=1
        report['Average Beta'] = df[inv_column_mapping['beta']].mean()
        doc.add_paragraph(f'Average Beta: {report["Average Beta"]}')

    if 'roe' in column_mapping.values():
        matched_columns_count+=1
        report['Average ROE'] = df[inv_column_mapping['roe']].mean()
        doc.add_paragraph(f'Average ROE: {report["Average ROE"]}')

    if 'news_sentiment' in column_mapping.values():
        matched_columns_count+=1
        sentiment_counts = df[inv_column_mapping['news_sentiment']].value_counts()
        doc.add_paragraph('News Sentiment Distribution:')
        for sentiment, count in sentiment_counts.items():
            doc.add_paragraph(f'{sentiment}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=sentiment_counts.index, y=sentiment_counts.values, ax=ax)
        ax.set_title('News Sentiment Distribution')
        save_plot_to_word(doc, fig, 'News Sentiment Distribution')

    if 'analyst_rating' in column_mapping.values():
        matched_columns_count+=1
        rating_counts = df[inv_column_mapping['analyst_rating']].value_counts()
        doc.add_paragraph('Analyst Rating Distribution:')
        for rating, count in rating_counts.items():
            doc.add_paragraph(f'{rating}: {count}')
        fig, ax = plt.subplots()
        sns.barplot(x=rating_counts.index, y=rating_counts.values, ax=ax)
        ax.set_title('Analyst Rating Distribution')
        save_plot_to_word(doc, fig, 'Analyst Rating Distribution')

    if matched_columns_count == 0:
        generate_plots_for_unmatched_columns(df, doc)
        
    return report

def generate_report(df, business_context, column_mapping, output_file='report.docx'):
    doc = Document()
    report_functions = {
        "Financial data": generate_financial_report,
        "Retail data": generate_retail_report,
        "Stock data": generate_stock_report,
        "College Data": generate_college_report,
        "Marketing data": generate_marketing_report,
        "Sales data": generate_sales_report,
        "Inventory data": generate_inventory_report,
        "Healthcare data": generate_healthcare_report,
    }
    inv_column_mapping = {v: k for k, v in column_mapping.items()}
    if business_context in report_functions:
        report_functions[business_context](df, column_mapping, inv_column_mapping, doc)
        doc.save(output_file)
        print(f'Report saved to {output_file}')
        return True
    else:
        print("Invalid business context")
        return False


@app.route('/generate_report', methods=['POST'])
def generate_report_route():
    try:
        data = request.get_json()
        business_context = data.get('business_context')
        column_mapping = data.get('column_mapping', {})
        user_id = data.get('user_id', 'unknown')
        
        if not business_context:
            return jsonify({'error': 'Missing business context'}), 400
            
        # Get the latest mapped dataset
        dataset_path = os.path.join(data_dir, 'mapped_dataset.csv')
        if not os.path.exists(dataset_path):
            # Fall back to the latest uploaded file
            files = [f for f in os.listdir(data_dir) if f.endswith(('.csv', '.xlsx'))]
            if not files:
                return jsonify({'error': 'No dataset file found'}), 400
            dataset_path = os.path.join(data_dir, files[-1])
        
        # Load dataset
        _, ext = os.path.splitext(dataset_path)
        if ext == '.csv':
            df = pd.read_csv(dataset_path)
        elif ext in ['.xls', '.xlsx']:
            df = pd.read_excel(dataset_path)
        else:
            return jsonify({"error": "Unsupported file format"}), 400
            
        # Generate report
        output_file = os.path.join(output_dir, f"{user_id}_report.docx")
        success = generate_report(df, business_context, column_mapping, output_file)
        
        if success:
            return jsonify({
                "status": "Report generated successfully",
                "report_path": output_file
            }), 200
        else:
            return jsonify({"error": "Failed to generate report. Invalid business context."}), 400
            
    except Exception as e:
        return jsonify({"error": f"Error generating report: {str(e)}"}), 500







#ChatBot Implementation




from flask import Flask, request, jsonify
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import base64
import io
import traceback
import google.generativeai as genai
from dotenv import load_dotenv
import os

# Load the environment variables from .env
load_dotenv()

# Get the key
gemini_key = os.getenv("GEMINI_API_KEY")
# Load Gemini API key
genai.configure(api_key=gemini_key)

# Load the latest dataset
def load_dataset():
    files = [f for f in os.listdir(data_dir) if f.endswith(('.csv', '.xlsx'))]
    if not files:
        return None
    path = os.path.join(data_dir, files[-1])
    return pd.read_csv(path) if path.endswith('.csv') else pd.read_excel(path)


def parse_query(query, dataset_context):
    """Generates Python code using Google GenAI for analyzing a dataset based on a user query."""

    prompt = f"""
You are an expert Python assistant specialized in data analysis using pandas, numpy, matplotlib, and scikit-learn.

The dataset is already loaded as a pandas DataFrame named `df`.
Dataset columns: {dataset_context}
User query: {query}

Your task:
- Generate Python code that accurately answers the query using appropriate data analysis or machine learning techniques.
- The final result must be stored in a variable called `result`.
- If visualization is needed, use matplotlib and ensure you call `plt.show()` to display the plot.
- The output must be only Python code, with no extra text or explanation.
- If the query contains typos in column names, intelligently map them to the closest correct column name from the dataset.
- If the user asks how changes to a specific column affect another (e.g., "what if I increase column X"), perform regression or classification analysis depending on the target column:
    - Carefully preprocess the data: handle missing values, convert categorical columns using label encoding, and ensure the model doesn't raise errors.
    - Perform the analysis on the original data, simulate the increase in the specified column, predict new values, and compare them with the original target.
    - If the target (`y`) is numeric, plot only the original vs predicted values using a scatter or line plot.
    - If the target (`y`) is categorical, generate an appropriate classification evaluation plot like a bar plot comparing class distributions.
- Always use thorough exception handling. Your code should:
    - Handle missing columns gracefully and inform the user to revise the query if needed.
    - Skip or impute null values properly.
    - Catch and handle any unexpected errors.
- If the query is casual and unrelated to analysis, respond casually but try to steer the conversation toward data analysis.

Respond with only valid and executable Python code, enclosed within triple backticks.
"""

    model = genai.GenerativeModel('gemini-2.0-flash')
    response = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            temperature=0.7,
            top_p=0.9,
            max_output_tokens=2000
        )
    )
    
    return response.text



def clean_generated_code(generated_text):
    if not generated_text:
        return "result = df.head()"
    code_lines = []
    in_code_block = False
    for line in generated_text.split('\n'):
        if line.strip().startswith('```python'):
            in_code_block = True
            continue
        if line.strip().startswith('```'):
            in_code_block = False
            continue
        if in_code_block or line.strip().startswith('df'):
            if "plt.show()" in line:
                continue  # 👈 skip plt.show()
            code_lines.append(line)
    cleaned_code = '\n'.join(code_lines).strip()
    if "result" not in cleaned_code:
        cleaned_code += "\nresult = df.head()"
    return cleaned_code

def map_column_names(generated_code,column_mapping):
    """ Ensures correct column names in generated code """
    for user_col, actual_col in column_mapping.items():
        generated_code = generated_code.replace(user_col, actual_col)
    return generated_code


def execute_generated_code(code, df):
    local_vars = {"df": df, "np": np, "pd": pd, "plt": plt}
    try:
        plt.clf()
        plt.close('all')

        # 👇 Remove plt.show() to prevent consuming the figure
        code = code.replace("plt.show()", "")

        exec(code, globals(), local_vars)
        result = local_vars.get("result")

        image_base64 = None
        if plt.get_fignums():
            plot_buffer = io.BytesIO()
            plt.savefig(plot_buffer, format='png')
            plot_buffer.seek(0)
            image_base64 = base64.b64encode(plot_buffer.read()).decode('utf-8')
            plt.close()

        return result.to_string() if isinstance(result, pd.DataFrame) else str(result), image_base64, None

    except Exception as e:
        return None, None, traceback.format_exc()



@app.route('/process_query', methods=['POST'])
def handle_process_query():
    try:
        data = request.get_json()
        query = data.get("query")
        column_mapping = data.get("column_mapping", {})
        if not query:
            return jsonify({"error": "Query is missing."}), 400

        df = load_dataset()
        if df is None:
            return jsonify({"error": "No dataset found."}), 400

        # Clean up column names and rename
        df.columns = df.columns.str.strip()
        df.rename(columns=column_mapping, inplace=True)
        print("DF Columns After Mapping:", df.columns.tolist())  # Check that placement_status exists

        dataset_context = ", ".join([f"{k} → {v}" for k, v in column_mapping.items()])
        
        generated_code = parse_query(query, dataset_context)
        generated_code = clean_generated_code(generated_code)
        # Use the mapping directly (do not invert it)
        generated_code = map_column_names(generated_code, column_mapping)

        print("Generated Code:\n", generated_code)

        result, image_data, error = execute_generated_code(generated_code, df)
        if error:
            return jsonify({"error": error}), 500

        # Return additional debug information in the JSON response
        return jsonify({
            "result": result,
            "image": image_data,
            "generated_code": generated_code,   # DEBUG: Show what code was generated
            "df_columns": df.columns.tolist(),    # DEBUG: Show DataFrame columns after mapping
            "message": "Query executed successfully."
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500




if __name__ == '__main__':
    app.run(debug=False, threaded=False)